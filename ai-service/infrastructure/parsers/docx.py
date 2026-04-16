from __future__ import annotations

import logging
from io import BytesIO
from typing import Any

from docx import Document

from core.ports.parser import ParsedDocument, ParsedDocumentSection, ParserPort

logger = logging.getLogger(__name__)


class DocxParser(ParserPort):
    CONTENT_TYPES = frozenset(
        [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document; charset=utf-8",
        ]
    )

    def supports(self, content_type: str) -> bool:
        return content_type.lower() in self.CONTENT_TYPES

    async def parse(self, content: bytes, metadata: dict[str, Any] | None = None) -> ParsedDocument:
        try:
            doc = Document(BytesIO(content))
        except Exception:
            logger.exception("Failed to open DOCX")
            raise

        sections: list[ParsedDocumentSection] = []
        current_heading: str | None = None
        current_content: list[str] = []
        section_order = 0

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            style_name = para.style.name if para.style else ""

            if style_name and any(
                heading in style_name.lower() for heading in ["heading", "title", "head"]
            ):
                if current_content:
                    sections.append(
                        ParsedDocumentSection(
                            heading=current_heading,
                            content="\n".join(current_content).strip(),
                            order=section_order,
                        )
                    )
                    section_order += 1
                    current_content = []

                current_heading = text
                current_content = []
            else:
                current_content.append(text)

        if current_content:
            sections.append(
                ParsedDocumentSection(
                    heading=current_heading,
                    content="\n".join(current_content).strip(),
                    order=section_order,
                )
            )

        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        if not sections:
            sections.append(
                ParsedDocumentSection(
                    heading=None,
                    content=full_text.strip(),
                    order=0,
                )
            )

        return ParsedDocument(
            document_text=full_text.strip(),
            sections=sections,
            metadata=metadata or {},
        )


__all__ = ["DocxParser"]
